# advanced_analyzer.py
"""
Advanced Detection and Risk Analysis Module
Sophisticated invisibility scoring and detection risk assessment
"""

import re
import unicodedata
import statistics
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
from enum import Enum
import docx
from pathlib import Path

class DetectionRisk(Enum):
    MINIMAL = "minimal"     # 0-0.2
    LOW = "low"            # 0.2-0.4  
    MEDIUM = "medium"      # 0.4-0.6
    HIGH = "high"          # 0.6-0.8
    CRITICAL = "critical"  # 0.8-1.0

@dataclass
class RiskAnalysis:
    overall_risk: float
    risk_level: DetectionRisk
    invisibility_score: float
    detection_patterns: List[str]
    recommendations: List[str]
    technique_breakdown: Dict[str, float]

@dataclass
class CharacterAnalysis:
    total_chars: int
    visible_changes: int
    invisible_changes: int
    unicode_substitutions: int
    zero_width_insertions: int
    suspicious_patterns: List[str]

class AdvancedAnalyzer:
    """Advanced analysis for steganographic modifications"""
    
    def __init__(self, logger=None):
        self.logger = logger
        
        # Known detection patterns used by plagiarism checkers
        self.detection_patterns = {
            'excessive_unicode': r'[А-Я]{3,}',  # Too many Cyrillic chars
            'zero_width_clusters': r'[\u200B-\u200D\uFEFF]{2,}',  # Clustered ZW chars
            'mixed_scripts': r'[a-zA-Z][А-Я]|[А-Я][a-zA-Z]',  # Script mixing
            'academic_keywords': r'(ВАВ|РENDAHULUAN|МETODE)',  # Substituted keywords
            'abnormal_spacing': r'\s{3,}',  # Excessive spacing
        }
        
        # Risk weights for different modification types
        self.risk_weights = {
            'unicode_density': 0.3,
            'zero_width_density': 0.2,
            'pattern_detection': 0.25,
            'script_mixing': 0.15,
            'modification_distribution': 0.1
        }
    
    def analyze_document_risk(
        self, 
        original_path: str, 
        modified_path: str
    ) -> RiskAnalysis:
        """Comprehensive risk analysis of modified document"""
        
        try:
            # Load documents
            original_doc = docx.Document(original_path)
            modified_doc = docx.Document(modified_path)
            
            # Extract text
            original_text = '\n'.join([p.text for p in original_doc.paragraphs])
            modified_text = '\n'.join([p.text for p in modified_doc.paragraphs])
            
            # Character-level analysis
            char_analysis = self._analyze_characters(original_text, modified_text)
            
            # Pattern detection
            detected_patterns = self._detect_suspicious_patterns(modified_text)
            
            # Risk calculation
            risk_scores = self._calculate_risk_components(char_analysis, detected_patterns)
            overall_risk = self._calculate_overall_risk(risk_scores)
            
            # Invisibility score
            invisibility_score = self._calculate_invisibility_score(char_analysis)
            
            # Generate recommendations
            recommendations = self._generate_recommendations(risk_scores, detected_patterns)
            
            return RiskAnalysis(
                overall_risk=overall_risk,
                risk_level=self._get_risk_level(overall_risk),
                invisibility_score=invisibility_score,
                detection_patterns=detected_patterns,
                recommendations=recommendations,
                technique_breakdown=risk_scores
            )
            
        except Exception as e:
            if self.logger:
                self.logger.error(f"Risk analysis failed: {e}")
            
            return RiskAnalysis(
                overall_risk=1.0,
                risk_level=DetectionRisk.CRITICAL,
                invisibility_score=0.0,
                detection_patterns=["Analysis failed"],
                recommendations=["Unable to analyze - check file integrity"],
                technique_breakdown={}
            )
    
    def _analyze_characters(self, original: str, modified: str) -> CharacterAnalysis:
        """Detailed character-level analysis"""
        
        visible_changes = 0
        invisible_changes = 0
        unicode_substitutions = 0
        zero_width_insertions = 0
        suspicious_patterns = []
        
        # Character by character comparison
        for i, (orig_char, mod_char) in enumerate(zip(original, modified)):
            if orig_char != mod_char:
                # Check if visually identical (Unicode substitution)
                if unicodedata.normalize('NFKC', orig_char) == unicodedata.normalize('NFKC', mod_char):
                    invisible_changes += 1
                    unicode_substitutions += 1
                else:
                    visible_changes += 1
        
        # Check for zero-width character insertions
        zw_chars = ['\u200B', '\u200C', '\u200D', '\uFEFF']
        for char in zw_chars:
            count = modified.count(char)
            if count > 0:
                zero_width_insertions += count
                invisible_changes += count
        
        # Detect suspicious patterns
        if unicode_substitutions > len(original) * 0.05:  # >5% substitutions
            suspicious_patterns.append("high_unicode_density")
        
        if zero_width_insertions > len(original) * 0.03:  # >3% ZW chars
            suspicious_patterns.append("excessive_zero_width")
        
        return CharacterAnalysis(
            total_chars=len(modified),
            visible_changes=visible_changes,
            invisible_changes=invisible_changes,
            unicode_substitutions=unicode_substitutions,
            zero_width_insertions=zero_width_insertions,
            suspicious_patterns=suspicious_patterns
        )
    
    def _detect_suspicious_patterns(self, text: str) -> List[str]:
        """Detect patterns that might trigger plagiarism detectors"""
        
        detected = []
        
        for pattern_name, regex in self.detection_patterns.items():
            matches = re.findall(regex, text)
            if matches:
                detected.append(f"{pattern_name}: {len(matches)} matches")
        
        return detected
    
    def _calculate_risk_components(
        self, 
        char_analysis: CharacterAnalysis, 
        patterns: List[str]
    ) -> Dict[str, float]:
        """Calculate individual risk component scores"""
        
        total_chars = char_analysis.total_chars
        
        # Unicode density risk
        unicode_density = char_analysis.unicode_substitutions / total_chars if total_chars > 0 else 0
        unicode_risk = min(1.0, unicode_density * 20)  # 5% = 1.0 risk
        
        # Zero-width density risk  
        zw_density = char_analysis.zero_width_insertions / total_chars if total_chars > 0 else 0
        zw_risk = min(1.0, zw_density * 33)  # 3% = 1.0 risk
        
        # Pattern detection risk
        pattern_risk = min(1.0, len(patterns) * 0.2)  # 5 patterns = 1.0 risk
        
        # Script mixing risk (Latin + Cyrillic in same word)
        script_mixing_risk = 0.5 if any('mixed_scripts' in p for p in patterns) else 0
        
        # Modification distribution (clustered vs distributed)
        distribution_risk = 0.3 if any('clusters' in p for p in patterns) else 0
        
        return {
            'unicode_density': unicode_risk,
            'zero_width_density': zw_risk,
            'pattern_detection': pattern_risk,
            'script_mixing': script_mixing_risk,
            'modification_distribution': distribution_risk
        }
    
    def _calculate_overall_risk(self, risk_scores: Dict[str, float]) -> float:
        """Calculate weighted overall risk score"""
        
        total_risk = 0
        for component, risk in risk_scores.items():
            weight = self.risk_weights.get(component, 0.1)
            total_risk += risk * weight
        
        return min(1.0, total_risk)
    
    def _calculate_invisibility_score(self, char_analysis: CharacterAnalysis) -> float:
        """Calculate how invisible the changes are (higher = better)"""
        
        total_changes = char_analysis.visible_changes + char_analysis.invisible_changes
        if total_changes == 0:
            return 1.0
        
        invisibility_ratio = char_analysis.invisible_changes / total_changes
        return invisibility_ratio
    
    def _get_risk_level(self, risk_score: float) -> DetectionRisk:
        """Convert risk score to risk level enum"""
        
        if risk_score < 0.2:
            return DetectionRisk.MINIMAL
        elif risk_score < 0.4:
            return DetectionRisk.LOW
        elif risk_score < 0.6:
            return DetectionRisk.MEDIUM
        elif risk_score < 0.8:
            return DetectionRisk.HIGH
        else:
            return DetectionRisk.CRITICAL
    
    def _generate_recommendations(
        self, 
        risk_scores: Dict[str, float], 
        patterns: List[str]
    ) -> List[str]:
        """Generate actionable recommendations to reduce detection risk"""
        
        recommendations = []
        
        # Unicode density recommendations
        if risk_scores.get('unicode_density', 0) > 0.6:
            recommendations.append("Reduce Unicode substitution rate (<3% of total text)")
        
        # Zero-width recommendations
        if risk_scores.get('zero_width_density', 0) > 0.5:
            recommendations.append("Decrease zero-width character insertion frequency")
        
        # Pattern recommendations
        if risk_scores.get('pattern_detection', 0) > 0.4:
            recommendations.append("Avoid consecutive Unicode substitutions in keywords")
        
        # Script mixing recommendations
        if risk_scores.get('script_mixing', 0) > 0.3:
            recommendations.append("Ensure consistent script usage within words")
        
        # Distribution recommendations
        if risk_scores.get('modification_distribution', 0) > 0.2:
            recommendations.append("Distribute modifications more evenly across document")
        
        # General recommendations
        if not recommendations:
            recommendations.append("Current modification level appears safe for most detectors")
        
        return recommendations
