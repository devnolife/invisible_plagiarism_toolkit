import os
from pathlib import Path
from invisible_manipulator import InvisibleManipulator

def test_unicode_mapping_normalized():
    manip = InvisibleManipulator(verbose=False)
    assert 'latin_to_cyrillic' in manip.unicode_mappings, "latin_to_cyrillic mapping missing after normalization"
    assert 'special_substitutions' in manip.unicode_mappings, "special_substitutions mapping missing after normalization"


def test_header_detection_simple(tmp_path):
    # Create minimal docx
    import docx
    file_path = tmp_path / 'sample.docx'
    d = docx.Document()
    d.add_paragraph('BAB I')
    d.add_paragraph('PENDAHULUAN')
    d.add_paragraph('A. Latar Belakang')
    d.add_paragraph('Konten biasa.')
    d.save(file_path)

    manip = InvisibleManipulator(verbose=False)
    analysis = manip.analyze_document_structure(str(file_path))
    headers = [h['text'] for h in analysis['headers']]
    assert 'BAB I' in headers
    assert 'PENDAHULUAN' in headers
    assert any('Latar Belakang' in h for h in headers)


def test_process_document(tmp_path):
    import docx
    file_path = tmp_path / 'proc.docx'
    d = docx.Document()
    d.add_paragraph('BAB I')
    d.add_paragraph('PENDAHULUAN')
    d.add_paragraph('Penelitian ini adalah contoh dan data uji.')
    d.save(file_path)

    manip = InvisibleManipulator(verbose=False)
    result = manip.apply_invisible_manipulation(str(file_path))
    assert result['output_file']
    assert os.path.exists(result['output_file'])
    assert result['stats']['processing_time'] >= 0
