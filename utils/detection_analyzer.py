"""
Lightweight detection/analyzer utilities.

Provides quick comparisons and invisibility scoring between two DOCX files.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Dict, Optional
import unicodedata

import docx


@dataclass
class InvisibilityReport:
	invisible_changes: int
	visible_changes: int
	total_chars_changed: int

	@property
	def invisibility_score(self) -> float:
		total = self.invisible_changes + self.visible_changes
		return (self.invisible_changes / total) if total else 1.0


def compare_docx_invisibility(original_path: str, modified_path: str) -> Optional[Dict]:
	"""Compare two DOCX docs and estimate invisibility metrics.

	Returns a dict compatible with main.verify_invisibility expectations.
	"""
	try:
		original_doc = docx.Document(original_path)
		modified_doc = docx.Document(modified_path)

		invisible_changes = 0
		visible_changes = 0
		total_chars_changed = 0

		for orig_para, mod_para in zip(original_doc.paragraphs, modified_doc.paragraphs):
			o = orig_para.text
			m = mod_para.text
			if o == m:
				continue

			o_vis = "".join(c for c in o if unicodedata.category(c) != "Cf")
			m_vis = "".join(c for c in m if unicodedata.category(c) != "Cf")

			if o_vis == m_vis:
				invisible_changes += 1
			else:
				visible_changes += 1

			total_chars_changed += abs(len(m) - len(o))

		return {
			"invisible_changes": invisible_changes,
			"visible_changes": visible_changes,
			"total_chars_changed": total_chars_changed,
		}
	except Exception:
		return None


__all__ = ["InvisibilityReport", "compare_docx_invisibility"]

