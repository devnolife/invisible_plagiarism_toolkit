# invisible_manipulator.py
"""
Invisible Manipulator Core - Main Engine
Sistem manipulasi dokumen menggunakan teknik steganografi dan karakter invisible
Fokus: Header manipulation, Unicode substitution, Invisible characters

Author: DevNoLife
Version: 1.0 - Steganography Edition
"""

import os
import json
import re
import random
import shutil
from datetime import datetime
from pathlib import Path
import logging
import docx
import unicodedata
from metadata_manipulator import MetadataManipulator, MetadataOptions
from unicode_steganography import UnicodeSteg
try:
    from utils.detection_analyzer import compare_docx_invisibility
except ImportError:
    # Fallback for when detection_analyzer is not available
    def compare_docx_invisibility(*args, **kwargs):
        return {"invisible_changes": 0, "visible_changes": 0, "total_chars_changed": 0}

class InvisibleManipulator:
    def __init__(self, config_file='config.json', verbose: bool = False):
        self.verbose = verbose
        self.logger = logging.getLogger(self.__class__.__name__)
        if not self.logger.handlers:
            handler = logging.StreamHandler()
            fmt = logging.Formatter('[%(levelname)s] %(name)s: %(message)s')
            handler.setFormatter(fmt)
            self.logger.addHandler(handler)
        self.logger.setLevel(logging.DEBUG if verbose else logging.INFO)

        self.logger.info("Initializing Invisible Manipulator (steganography engine)")

        # Load configuration
        self.config = self.load_config(config_file)
        self.validate_config(self.config)

        # Load data files
        self.unicode_mappings = self.load_data_file('data/unicode_mappings.json')
        self.invisible_chars = self.load_data_file('data/invisible_chars.json')
        self.header_patterns = self.load_data_file('data/header_patterns.json')

        # Fallback: initialize UnicodeSteg if mapping keys mismatch expected usage
        self.steg = UnicodeSteg()
        if 'latin_cyrillic' in self.steg.mappings and 'latin_to_cyrillic' not in self.unicode_mappings:
            # Normalize structure so downstream code can use unified key names
            self.unicode_mappings['latin_to_cyrillic'] = self.steg.mappings['latin_cyrillic']
        if 'indonesian_academic' in self.steg.mappings and 'special_substitutions' not in self.unicode_mappings:
            self.unicode_mappings['special_substitutions'] = self.steg.mappings['indonesian_academic']

        # Statistics tracking (per document basis later)
        self.reset_stats()

        self.logger.debug("Configuration loaded and mappings normalized")

    def reset_stats(self):
        self.stats = {
            'total_documents': 0,
            'headers_modified': 0,
            'chars_substituted': 0,
            'invisible_chars_inserted': 0,
            'metadata_modified': 0,
            'processing_time': 0
        }
    
    def load_config(self, config_file):
        """Load configuration from JSON file"""
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
            print(f"üìã Configuration loaded from {config_file}")
            return config
        except FileNotFoundError:
            print(f"‚ö†Ô∏è Config file {config_file} not found, using defaults")
            return self.get_default_config()

    def validate_config(self, config: dict):
        """Basic schema validation and normalization."""
        required_root = ['invisible_techniques', 'safety_settings']
        for key in required_root:
            if key not in config:
                self.logger.warning("Config missing root key '%s' - inserting defaults", key)
                config[key] = self.get_default_config()[key]
        try:
            z = config['invisible_techniques'].get('zero_width_chars', {})
            if 'insertion_rate' in z:
                z['insertion_rate'] = max(0.0, min(0.2, float(z['insertion_rate'])))
            u = config['invisible_techniques'].get('unicode_substitution', {})
            if 'substitution_rate' in u:
                u['substitution_rate'] = max(0.0, min(0.2, float(u['substitution_rate'])))
        except Exception as e:
            self.logger.warning("Rate normalization failed: %s", e)
    
    def load_data_file(self, file_path):
        """Load data files (mappings, patterns, etc.)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data
        except FileNotFoundError:
            print(f"‚ö†Ô∏è Data file {file_path} not found, using defaults")
            return {}
    
    def get_default_config(self):
        """Default configuration if config file is missing"""
        return {
            "invisible_techniques": {
                "zero_width_chars": {
                    "enabled": True,
                    "insertion_rate": 0.05,
                    "target_locations": ["headers", "after_punctuation"]
                },
                "unicode_substitution": {
                    "enabled": True,
                    "substitution_rate": 0.03
                },
                "metadata_manipulation": {
                    "enabled": True
                }
            },
            "safety_settings": {
                "preserve_readability": True,
                "backup_original": True,
                "max_changes_per_paragraph": 5
            }
        }
    
    def analyze_document_structure(self, doc_path):
        """Analyze document to identify headers, key sections, etc."""
        self.logger.debug("Analyzing document structure ...")

        doc = docx.Document(doc_path)
        analysis = {
            'total_paragraphs': len(doc.paragraphs),
            'headers': [],
            'key_sections': [],
            'citations': [],
            'metadata': {}
        }

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            if not text:
                continue
            
            # Detect headers
            if self.is_header(text, paragraph):
                analysis['headers'].append({
                    'index': i,
                    'text': text,
                    'type': self.classify_header(text),
                    'priority': self.get_header_priority(text)
                })
            
            # Detect key sections
            if self.is_key_section(text):
                analysis['key_sections'].append({
                    'index': i,
                    'text': text,
                    'section_type': self.classify_section(text)
                })
            
            # Detect citations (simple pattern)
            if self.contains_citation(text):
                analysis['citations'].append({
                    'index': i,
                    'text': text[:100] + '...' if len(text) > 100 else text
                })
        self.logger.debug(
            "Analysis: paragraphs=%s headers=%s key_sections=%s citations=%s",
            analysis['total_paragraphs'], len(analysis['headers']), len(analysis['key_sections']), len(analysis['citations'])
        )
        return analysis
    
    def is_header(self, text, paragraph):
        """Determine if text is a header"""
        # Check formatting clues
        if hasattr(paragraph, 'style') and paragraph.style:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name or 'title' in style_name:
                return True
        
        # Check text patterns
        text_upper = text.upper()
        
        # Common header patterns
        header_indicators = [
            r'^BAB\s+[IVX]+',           # BAB I, BAB II, etc.
            r'^CHAPTER\s+\d+',         # CHAPTER 1, etc.
            r'^[A-Z]\.\s+[A-Z]',       # A. SOMETHING
            r'^PENDAHULUAN$',          # Exact matches
            r'^METODE\s+PENELITIAN$',
            r'^HASIL\s+DAN\s+PEMBAHASAN$',
            r'^KESIMPULAN$',
            r'^DAFTAR\s+PUSTAKA$'
        ]
        
        for pattern in header_indicators:
            if re.match(pattern, text_upper):
                return True
        
        # Additional checks
        conditions = [
            len(text.split()) <= 6,     # Short
            text.isupper(),             # All caps
            len(text) < 50,             # Not too long
            not text.endswith('.')      # Usually no period
        ]
        
        return sum(conditions) >= 2
    
    def classify_header(self, text):
        """Classify header type"""
        text_upper = text.upper()
        
        if re.match(r'^BAB\s+[IVX]+', text_upper):
            return 'chapter'
        elif text_upper in ['PENDAHULUAN', 'METODE PENELITIAN', 'HASIL DAN PEMBAHASAN', 'KESIMPULAN']:
            return 'major_section'
        elif re.match(r'^[A-Z]\.\s+', text_upper):
            return 'subsection'
        else:
            return 'minor_header'
    
    def get_header_priority(self, text):
        """Get manipulation priority for header"""
        text_upper = text.upper()
        
        if 'BAB' in text_upper or text_upper in ['PENDAHULUAN', 'KESIMPULAN']:
            return 'highest'
        elif text_upper in ['METODE PENELITIAN', 'HASIL DAN PEMBAHASAN']:
            return 'high'
        elif re.match(r'^[A-Z]\.\s+', text_upper):
            return 'medium'
        else:
            return 'low'
    
    def is_key_section(self, text):
        """Identify key sections that need attention"""
        key_phrases = [
            'latar belakang', 'rumusan masalah', 'tujuan penelitian',
            'metode penelitian', 'hasil penelitian', 'analisis data',
            'kesimpulan', 'saran', 'daftar pustaka'
        ]
        
        text_lower = text.lower()
        return any(phrase in text_lower for phrase in key_phrases)
    
    def classify_section(self, text):
        """Classify section type"""
        text_lower = text.lower()
        
        if 'latar belakang' in text_lower:
            return 'background'
        elif 'rumusan masalah' in text_lower:
            return 'problem_statement'
        elif 'tujuan' in text_lower:
            return 'objectives'
        elif 'metode' in text_lower:
            return 'methodology'
        elif 'hasil' in text_lower:
            return 'results'
        elif 'kesimpulan' in text_lower:
            return 'conclusion'
        else:
            return 'general'
    
    def contains_citation(self, text):
        """Simple citation detection"""
        citation_patterns = [
            r'\(\d{4}\)',               # (2020)
            r'\([^)]*\d{4}[^)]*\)',    # (Smith, 2020)
            r'et\s+al\.',              # et al.
            r'\w+\s+\(\d{4}\)'         # Author (2020)
        ]
        
        return any(re.search(pattern, text) for pattern in citation_patterns)
    
    def apply_invisible_manipulation(self, doc_path, output_path=None, dry_run: bool = False):
        """Apply invisible manipulation techniques.

        dry_run: simulate; no file writes/backups.
        """
        self.logger.info("Starting invisible manipulation: %s", doc_path)

        start_time = datetime.now()

        # Create backup
        backup_path = None
        if not dry_run and self.config['safety_settings']['backup_original']:
            backup_path = self.create_backup(doc_path)
            self.logger.debug("Backup created: %s", backup_path)

        # Analyze document
        analysis = self.analyze_document_structure(doc_path)

        # Load document
        doc = docx.Document(doc_path)

        # Apply techniques based on priority
        self.apply_header_manipulation(doc, analysis)
        self.apply_unicode_substitution(doc, analysis)
        self.apply_invisible_characters(doc, analysis)
        self.apply_metadata_manipulation(doc)

        # Save processed document
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = Path(doc_path).stem
            output_path = f"output/processed_documents/{base_name}_invisible_{timestamp}.docx"
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        if not dry_run:
            doc.save(output_path)
        
        # Calculate statistics
        end_time = datetime.now()
        self.stats['processing_time'] = (end_time - start_time).total_seconds()
        self.stats['total_documents'] += 1

        if dry_run:
            self.logger.info("Dry-run completed (no file written)")
        else:
            self.logger.info("Invisible manipulation completed -> %s", output_path)
        self.print_manipulation_stats()

        return {
            'input_file': doc_path,
            'output_file': None if dry_run else output_path,
            'dry_run': dry_run,
            'backup_file': backup_path if (not dry_run and self.config['safety_settings']['backup_original']) else None,
            'analysis': analysis,
            'stats': self.stats.copy(),
            'processing_time': self.stats['processing_time']
        }
    
    def apply_header_manipulation(self, doc, analysis):
        """Apply manipulation specifically to headers"""
        self.logger.debug("Manipulating headers ...")
        
        if not self.config['invisible_techniques']['unicode_substitution']['enabled']:
            return
        
        headers = analysis['headers']
        high_priority_headers = [h for h in headers if h['priority'] in ['highest', 'high']]
        
        for header_info in high_priority_headers:
            paragraph = doc.paragraphs[header_info['index']]
            original_text = paragraph.text
            
            # Apply Unicode substitution for headers
            manipulated_text = self.apply_unicode_substitution_to_text(original_text)
            
            if manipulated_text != original_text:
                paragraph.text = manipulated_text
                self.stats['headers_modified'] += 1
                
                self.logger.debug("Header modified: %s -> %s", original_text, manipulated_text)
    
    def apply_unicode_substitution_to_text(self, text):
        """Apply Unicode character substitution"""
        if not self.unicode_mappings:
            return text
        
        substitution_rate = self.config['invisible_techniques']['unicode_substitution']['substitution_rate']
        latin_to_cyrillic = self.unicode_mappings.get('latin_to_cyrillic', {})
        special_subs = self.unicode_mappings.get('special_substitutions', {})
        
        result = text
        
        # Apply special substitutions first (whole words)
        for original, replacement in special_subs.items():
            if original in result:
                if random.random() < substitution_rate * 3:  # Higher rate for special words
                    result = result.replace(original, replacement)
                    self.stats['chars_substituted'] += len(original)
        
        # Apply individual character substitutions
        new_result = ""
        for char in result:
            if char in latin_to_cyrillic and random.random() < substitution_rate:
                new_result += latin_to_cyrillic[char]
                self.stats['chars_substituted'] += 1
            else:
                new_result += char
        
        return new_result
    
    def apply_unicode_substitution(self, doc, analysis):
        """Apply Unicode substitution to key sections"""
        self.logger.debug("Applying Unicode substitutions to key sections ...")

        if not self.config['invisible_techniques']['unicode_substitution']['enabled']:
            return

        key_sections = analysis['key_sections']

        for section_info in key_sections:
            paragraph = doc.paragraphs[section_info['index']]
            original_text = paragraph.text

            # Apply substitution with lower rate for content
            manipulated_text = self.apply_unicode_substitution_to_text(original_text)

            if manipulated_text != original_text:
                paragraph.text = manipulated_text
    
    def apply_invisible_characters(self, doc, analysis):
        """Insert invisible characters strategically"""
        self.logger.debug("Inserting invisible characters ...")

        if not self.config['invisible_techniques']['zero_width_chars']['enabled']:
            return

        if not self.invisible_chars:
            return

        insertion_rate = self.config['invisible_techniques']['zero_width_chars']['insertion_rate']
        zero_width_chars = list(self.invisible_chars.get('zero_width', {}).values())

        if not zero_width_chars:
            return

        # Focus on headers and key sections
        target_paragraphs = []

        # Add headers
        for header_info in analysis['headers']:
            target_paragraphs.append(header_info['index'])

        # Add key sections
        for section_info in analysis['key_sections']:
            target_paragraphs.append(section_info['index'])

        for para_index in target_paragraphs:
            if para_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[para_index]
                original_text = paragraph.text

                # Insert invisible characters
                new_text = self.insert_invisible_chars(original_text, zero_width_chars, insertion_rate)
                
                if new_text != original_text:
                    paragraph.text = new_text
                    self.stats['invisible_chars_inserted'] += new_text.count('\u200B') + new_text.count('\u200C')
    
    def insert_invisible_chars(self, text, invisible_chars, insertion_rate):
        """Insert invisible characters into text"""
        result = ""
        
        max_consecutive = self.invisible_chars.get('detection_avoidance', {}).get('max_consecutive', 2)
        consecutive = 0
        for i, char in enumerate(text):
            result += char

            # Rate limiting per paragraph
            if random.random() < insertion_rate and char in '.,:;! ':
                if consecutive < max_consecutive:
                    invisible_char = random.choice(invisible_chars)
                    result += invisible_char
                    consecutive += 1
                else:
                    consecutive = 0
        
        return result
    
    def apply_metadata_manipulation(self, doc):
        """Manipulate document metadata"""
        self.logger.debug("Manipulating metadata ...")
        if not self.config['invisible_techniques']['metadata_manipulation']['enabled']:
            return

        try:
            options = MetadataOptions(
                modify_properties=self.config['invisible_techniques']['metadata_manipulation'].get('modify_properties', True),
                add_invisible_content=self.config['invisible_techniques']['metadata_manipulation'].get('add_invisible_content', True),
            )
            manip = MetadataManipulator(options)
            changes = manip.apply(doc)
            self.stats['metadata_modified'] += int(changes > 0)
        except Exception as e:
            self.logger.warning("Could not modify metadata: %s", e)
    
    def create_backup(self, file_path):
        """Create backup of original file"""
        backup_dir = Path("backup")
        backup_dir.mkdir(exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_name = Path(file_path).stem
        extension = Path(file_path).suffix
        
        backup_path = backup_dir / f"{original_name}_backup_{timestamp}{extension}"
        
        shutil.copy2(file_path, backup_path)
        return str(backup_path)
    
    def print_manipulation_stats(self):
        """Print manipulation statistics"""
        self.logger.info(
            "Stats: docs=%s headers=%s chars_sub=%s inv_chars=%s metadata=%s time=%.2fs",
            self.stats['total_documents'], self.stats['headers_modified'], self.stats['chars_substituted'],
            self.stats['invisible_chars_inserted'], self.stats['metadata_modified'], self.stats['processing_time']
        )
    
    def verify_invisibility(self, original_path, modified_path):
        """Verify that changes are visually invisible (delegates to compare_docx_invisibility)."""
        self.logger.debug("Verifying invisibility of changes (unified analyzer)...")
        try:
            differences = compare_docx_invisibility(original_path, modified_path)
            if differences is None:
                return None
            # Enrich with score for convenience
            total_pairs = differences['invisible_changes'] + differences['visible_changes']
            differences['invisibility_score'] = (
                differences['invisible_changes'] / total_pairs if total_pairs else 1.0
            )
            self.logger.debug(
                "Verification result: invisible=%s visible=%s total_char_delta=%s score=%.1f%%",
                differences['invisible_changes'], differences['visible_changes'], differences['total_chars_changed'], differences['invisibility_score'] * 100
            )
            return differences
        except Exception as e:
            self.logger.error("Could not verify invisibility: %s", e)
            return None


def create_sample_config():
    """Create sample configuration file"""
    config = {
        "invisible_techniques": {
            "zero_width_chars": {
                "enabled": True,
                "chars": ["\u200B", "\u200C", "\u200D", "\uFEFF"],
                "insertion_rate": 0.05,
                "target_locations": ["headers", "after_punctuation", "between_words"]
            },
            "unicode_substitution": {
                "enabled": True,
                "confidence_level": "high",
                "target_chars": ["a", "e", "o", "p", "c", "x", "y"],
                "substitution_rate": 0.03
            },
            "spacing_manipulation": {
                "enabled": False,
                "micro_adjustments": True,
                "spacing_variance": 0.1
            },
            "metadata_manipulation": {
                "enabled": True,
                "modify_properties": True,
                "add_invisible_content": True
            }
        },
        "safety_settings": {
            "preserve_readability": True,
            "maintain_formatting": True,
            "backup_original": True,
            "max_changes_per_paragraph": 5,
            "avoid_obvious_patterns": True
        },
        "target_elements": {
            "headers": {
                "priority": "highest",
                "techniques": ["unicode_substitution", "invisible_chars"]
            },
            "first_sentences": {
                "priority": "high",
                "techniques": ["spacing_manipulation"]
            },
            "citations": {
                "priority": "medium",
                "techniques": ["invisible_chars"]
            },
            "conclusion_paragraphs": {
                "priority": "high",
                "techniques": ["unicode_substitution"]
            }
        }
    }
    
    with open('config.json', 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    
    print("‚úÖ Sample config.json created")


def main():
    """Main function for testing"""
    print("üîÆ INVISIBLE MANIPULATOR - TEST MODE")
    print("=" * 50)
    
    # Create sample config if it doesn't exist
    if not os.path.exists('config.json'):
        create_sample_config()
    
    # Initialize manipulator
    manipulator = InvisibleManipulator()
    
    # Test with sample document
    input_file = "input/sample_thesis.docx"
    
    if os.path.exists(input_file):
        print(f"üéØ Processing: {input_file}")
        
        result = manipulator.apply_invisible_manipulation(input_file)
        
        if result:
            print("\nüéâ Processing completed successfully!")
            print(f"üì• Input: {result['input_file']}")
            print(f"üì§ Output: {result['output_file']}")
            
            if result['backup_file']:
                print(f"üíæ Backup: {result['backup_file']}")
            
            # Verify invisibility
            if result['backup_file']:
                manipulator.verify_invisibility(result['backup_file'], result['output_file'])
        
    else:
        print(f"‚ùå Input file not found: {input_file}")
        print("üí° Please create the input directory and add a sample document")


if __name__ == "__main__":
    main()
