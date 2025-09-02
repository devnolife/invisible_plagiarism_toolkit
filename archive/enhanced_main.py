# enhanced_main.py
"""
Enhanced Main Module dengan Advanced Bypass Engine
Versi yang lebih canggih untuk mengatasi detector plagiarisme modern
"""

import os
import sys
import json
import argparse
import random
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
import docx

# Import modules
from invisible_manipulator import InvisibleManipulator
from advanced_bypass_engine import AdvancedBypassEngine, create_ultra_aggressive_config
from advanced_analyzer import AdvancedAnalyzer
from logger_config import get_project_logger

class EnhancedInvisibleProcessor:
    """Enhanced processor dengan multiple bypass engines"""
    
    def __init__(self, config_file='config.json', aggression_level='extreme', verbose=False):
        self.logger = get_project_logger(verbose)
        self.aggression_level = aggression_level
        
        # Initialize engines
        self.traditional_engine = InvisibleManipulator(config_file, verbose)
        self.advanced_engine = AdvancedBypassEngine(aggression_level)
        self.analyzer = AdvancedAnalyzer(self.logger)
        
        self.logger.info(f"Enhanced processor initialized (Level: {aggression_level})")
    
    def process_document_advanced(self, file_path: str, target_detector: str = "turnitin") -> Dict:
        """Process document dengan advanced bypass techniques"""
        
        start_time = time.time()
        self.logger.info(f"Processing: {Path(file_path).name}")
        
        try:
            # 1. Load document
            doc = docx.Document(file_path)
            original_text = '\n'.join([p.text for p in doc.paragraphs])
            
            if not original_text.strip():
                raise ValueError("Document appears to be empty")
            
            # 2. Apply advanced bypass techniques
            self.logger.info("Applying advanced bypass techniques...")
            bypass_result = self.advanced_engine.apply_comprehensive_bypass(
                original_text, 
                target_detector=target_detector
            )
            
            # 3. Apply traditional techniques untuk double layer protection
            self.logger.info("Applying traditional steganography layer...")
            
            # Create temporary file dengan hasil bypass
            temp_doc = docx.Document()
            for line in bypass_result.modified_text.split('\n'):
                temp_doc.add_paragraph(line)
            
            temp_path = f"temp_advanced_{int(time.time())}.docx"
            temp_doc.save(temp_path)
            
            try:
                traditional_result = self.traditional_engine.apply_invisible_manipulation(temp_path)
            finally:
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            # 4. Final quality check
            final_doc_path = traditional_result['output_file']
            
            # 5. Comprehensive risk analysis
            self.logger.info("Performing risk analysis...")
            risk_analysis = self.analyzer.analyze_document_risk(file_path, final_doc_path)
            
            # 6. Create comprehensive report
            processing_time = time.time() - start_time
            
            result = {
                'original_file': file_path,
                'output_file': final_doc_path,
                'processing_time': processing_time,
                'bypass_techniques_used': bypass_result.bypass_techniques_used,
                'changes_summary': {
                    'advanced_changes': bypass_result.changes_made,
                    'traditional_changes': traditional_result['stats']['chars_substituted'] + 
                                         traditional_result['stats']['invisible_chars_inserted'],
                    'total_changes': bypass_result.changes_made + traditional_result['stats']['chars_substituted'] + 
                                   traditional_result['stats']['invisible_chars_inserted']
                },
                'invisibility_metrics': {
                    'advanced_invisibility': bypass_result.invisibility_score,
                    'overall_invisibility': risk_analysis.invisibility_score,
                    'detection_risk': risk_analysis.overall_risk,
                    'risk_level': risk_analysis.risk_level.value
                },
                'risk_analysis': {
                    'overall_risk': risk_analysis.overall_risk,
                    'risk_level': risk_analysis.risk_level.value,
                    'detection_patterns': risk_analysis.detection_patterns,
                    'recommendations': risk_analysis.recommendations,
                    'technique_breakdown': risk_analysis.technique_breakdown
                },
                'quality_assessment': self._assess_quality(original_text, bypass_result.modified_text),
                'success': True
            }
            
            # 7. Save detailed report
            self._save_comprehensive_report(result)
            
            self.logger.info(f"✅ Processing completed successfully!")
            self.logger.info(f"⚡ Total changes: {result['changes_summary']['total_changes']}")
            self.logger.info(f"🎯 Invisibility score: {result['invisibility_metrics']['overall_invisibility']:.1%}")
            self.logger.info(f"⚠️ Detection risk: {result['invisibility_metrics']['risk_level']}")
            
            return result
            
        except Exception as e:
            self.logger.error(f"Processing failed: {e}")
            raise
    
    def _assess_quality(self, original: str, modified: str) -> Dict:
        """Assess quality of modifications"""
        
        # Basic quality metrics
        original_words = len(original.split())
        modified_words = len(modified.split())
        
        return {
            'readability_preserved': abs(original_words - modified_words) / original_words < 0.1,
            'length_change_percent': (modified_words - original_words) / original_words * 100,
            'character_count_original': len(original),
            'character_count_modified': len(modified),
            'estimated_reading_impact': 'minimal' if abs(original_words - modified_words) < original_words * 0.05 else 'moderate'
        }
    
    def _save_comprehensive_report(self, result: Dict):
        """Save comprehensive processing report"""
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"output/analysis_reports/enhanced_report_{timestamp}.json"
        
        # Ensure directory exists
        Path(report_file).parent.mkdir(parents=True, exist_ok=True)
        
        # Create detailed report
        report = {
            'metadata': {
                'timestamp': datetime.now().isoformat(),
                'processor_version': 'Enhanced v2.0',
                'aggression_level': self.aggression_level
            },
            'processing_result': result,
            'recommendations': self._generate_user_recommendations(result),
            'next_steps': self._suggest_next_steps(result)
        }
        
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        
        self.logger.info(f"📊 Comprehensive report saved: {report_file}")
    
    def _generate_user_recommendations(self, result: Dict) -> List[str]:
        """Generate actionable recommendations for user"""
        
        recommendations = []
        risk_level = result['invisibility_metrics']['risk_level']
        
        if risk_level in ['minimal', 'low']:
            recommendations.append("✅ Document appears well-protected against most detectors")
            recommendations.append("💡 Consider testing with actual detector if possible")
        elif risk_level == 'medium':
            recommendations.append("⚠️ Moderate detection risk - consider additional processing")
            recommendations.append("🔄 Try re-processing with 'extreme' aggression level")
        else:
            recommendations.append("🚨 High detection risk detected!")
            recommendations.append("📝 Consider manual paraphrasing of key sections")
            recommendations.append("🔄 Re-process with different techniques")
        
        # Specific technique recommendations
        total_changes = result['changes_summary']['total_changes']
        if total_changes < 100:
            recommendations.append("📈 Consider increasing modification rate for better protection")
        
        return recommendations
    
    def _suggest_next_steps(self, result: Dict) -> List[str]:
        """Suggest next steps based on results"""
        
        steps = []
        
        steps.append("1. Review the generated document for readability")
        steps.append("2. Test with your target plagiarism checker if available")
        steps.append("3. If detected, try re-processing with higher aggression")
        steps.append("4. Consider manual review of flagged sections")
        steps.append("5. Keep original backup for reference")
        
        return steps

def create_enhanced_parser() -> argparse.ArgumentParser:
    """Create enhanced argument parser"""
    
    parser = argparse.ArgumentParser(
        description="🚀 Enhanced Invisible Plagiarism Toolkit - Advanced Bypass Engine",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    # Input options
    parser.add_argument("--file", "-f", required=True, help="Document to process")
    parser.add_argument("--output", "-o", help="Output directory")
    
    # Processing options
    parser.add_argument(
        "--aggression", "-a",
        choices=["moderate", "high", "extreme"],
        default="extreme",
        help="Bypass aggression level (default: extreme)"
    )
    
    parser.add_argument(
        "--detector", "-d",
        choices=["turnitin", "copyscape", "plagscan", "grammarly", "general"],
        default="turnitin",
        help="Target plagiarism detector (default: turnitin)"
    )
    
    parser.add_argument(
        "--layers", "-l",
        type=int,
        default=2,
        help="Number of bypass layers to apply (default: 2)"
    )
    
    # Analysis options
    parser.add_argument("--analyze-only", action="store_true", help="Only analyze, don't modify")
    parser.add_argument("--quality-check", action="store_true", help="Perform quality assessment")
    
    # Utility options
    parser.add_argument("--verbose", "-v", action="store_true", help="Verbose output")
    parser.add_argument("--seed", type=int, help="Random seed for reproducibility")
    
    return parser

def main():
    """Enhanced main function"""
    
    parser = create_enhanced_parser()
    args = parser.parse_args()
    
    # Set random seed
    if args.seed:
        random.seed(args.seed)
        print(f"🎲 Random seed set: {args.seed}")
    
    # Initialize enhanced processor
    processor = EnhancedInvisibleProcessor(
        aggression_level=args.aggression,
        verbose=args.verbose
    )
    
    print("🚀 Enhanced Invisible Plagiarism Toolkit v2.0")
    print(f"🎯 Target detector: {args.detector}")
    print(f"⚡ Aggression level: {args.aggression}")
    print(f"📄 Processing: {Path(args.file).name}")
    print("-" * 50)
    
    try:
        if args.analyze_only:
            # Analyze existing document
            print("🔍 Analyzing document...")
            # Add analysis-only functionality here
        else:
            # Full processing
            result = processor.process_document_advanced(
                args.file, 
                target_detector=args.detector
            )
            
            print("\n📊 PROCESSING RESULTS:")
            print(f"✅ Output file: {result['output_file']}")
            print(f"⚡ Total changes: {result['changes_summary']['total_changes']}")
            print(f"🎯 Invisibility: {result['invisibility_metrics']['overall_invisibility']:.1%}")
            print(f"⚠️  Risk level: {result['invisibility_metrics']['risk_level'].upper()}")
            print(f"⏱️  Processing time: {result['processing_time']:.2f}s")
            
            print("\n💡 RECOMMENDATIONS:")
            for rec in result['risk_analysis']['recommendations']:
                print(f"  • {rec}")
            
            print(f"\n📋 Detailed report available in output/analysis_reports/")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
